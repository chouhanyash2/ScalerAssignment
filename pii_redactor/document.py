"""
document.py
===========
Read and write .docx files with PII redaction, preserving run-level formatting.

Architecture:
  - DocxRedactor.redact(input_path, output_path) is the public entry point.
  - Internally processes: body paragraphs → tables (recursive) → headers/footers.
  - Each paragraph is analysed as a full string, then replacements are mapped
    back to individual runs to preserve bold/italic/colour/font.

Run-mapping strategy:
  We maintain a character-offset map of which run owns which character in the
  paragraph's concatenated text. Replacements are applied run-by-run, splitting
  runs if a PII span spans a run boundary.

Stats are collected in RedactionStats and logged at the end.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from presidio_analyzer import AnalyzerEngine, RecognizerResult
from presidio_anonymizer import AnonymizerEngine

from pii_redactor.consistency import ConsistencyMapper
from pii_redactor.engine import analyze_text
from pii_redactor.operators import FakerOperators

logger = logging.getLogger(__name__)


# ─── Stats ────────────────────────────────────────────────────────────────────

@dataclass
class RedactionStats:
    entity_counts: Dict[str, int] = field(default_factory=dict)
    detections: List[Dict[str, object]] = field(default_factory=list)

    def record(self, entity_type: str, n: int = 1) -> None:
        self.entity_counts[entity_type] = self.entity_counts.get(entity_type, 0) + n

    def record_detection(
        self,
        *,
        entity_type: str,
        start: int,
        end: int,
        text: str,
        score: float,
        source: str,
    ) -> None:
        self.record(entity_type)
        self.detections.append({
            "entity_type": entity_type,
            "start": start,
            "end": end,
            "text": text,
            "score": round(score, 4),
            "source": source,
        })

    @property
    def total(self) -> int:
        return sum(self.entity_counts.values())

    def report(self) -> str:
        lines = ["\n── Redaction Summary ─────────────────────────────────────────"]
        for etype, count in sorted(self.entity_counts.items(), key=lambda x: -x[1]):
            lines.append(f"  {etype:<25} {count:>6} replacements")
        lines.append(f"  {'─'*39}")
        lines.append(f"  {'TOTAL':<25} {self.total:>6} replacements")
        lines.append("──────────────────────────────────────────────────────────────")
        return "\n".join(lines)


# ─── Main Redactor ────────────────────────────────────────────────────────────

class DocxRedactor:
    """
    End-to-end redactor for Word documents.

    Usage:
        redactor = DocxRedactor(analyzer, anonymizer, faker_ops, mapper)
        stats = redactor.redact(input_path, output_path)
    """

    def __init__(
        self,
        analyzer: AnalyzerEngine,
        anonymizer: AnonymizerEngine,
        faker_operators: FakerOperators,
        consistency_mapper: ConsistencyMapper,
        language: str = "en",
    ) -> None:
        self._analyzer = analyzer
        self._anonymizer = anonymizer
        self._operators = faker_operators.operators
        self._mapper = consistency_mapper
        self._language = language
        self._stats = RedactionStats()
        self._seen_table_cells: set[object] = set()

    # ── Public entry point ────────────────────────────────────────────────────

    def redact(self, input_path: Path, output_path: Path) -> RedactionStats:
        logger.info("Opening: %s", input_path)
        doc = Document(str(input_path))
        self._stats = RedactionStats()
        self._seen_table_cells = set()

        # 1. Body paragraphs
        for index, para in enumerate(doc.paragraphs):
            self._redact_paragraph(para, source=f"body_para_{index}")

        # 2. Tables (including nested)
        for table_index, table in enumerate(doc.tables):
            self._redact_table(table, source_prefix=f"table_{table_index}")

        # 3. Headers & footers for all sections
        for section in doc.sections:
            for hf in (
                section.header,
                section.even_page_header,
                section.first_page_header,
                section.footer,
                section.even_page_footer,
                section.first_page_footer,
            ):
                if hf is not None:
                    for para_index, para in enumerate(hf.paragraphs):
                        self._redact_paragraph(
                            para,
                            source=f"section_{section.start_type}_header_footer_para_{para_index}",
                        )

        doc.save(str(output_path))
        logger.info("Saved: %s", output_path)
        self._mapper.log_summary(verbose=False)
        logger.info(self._stats.report())
        return self._stats

    # ── Paragraph processing ──────────────────────────────────────────────────

    def _redact_paragraph(self, para: Paragraph, source: str) -> None:
        full_text = para.text
        if not full_text.strip():
            return

        results = analyze_text(self._analyzer, full_text, language=self._language)
        if not results:
            return

        # Build replacement map: {original_span: fake_value}
        replacement_map = self._build_replacement_map(full_text, results, source)

        # Apply replacements preserving run formatting
        self._apply_replacements_to_runs(para, full_text, replacement_map)

    def _build_replacement_map(
        self,
        text: str,
        results: List[RecognizerResult],
        source: str,
    ) -> Dict[Tuple[int, int], str]:
        """
        Returns {(start, end): fake_value} for each detected span.
        The ConsistencyMapper ensures the same original always gives the same fake.
        """
        replacement_map: Dict[Tuple[int, int], str] = {}

        for res in sorted(results, key=lambda r: r.start, reverse=True):
            original = text[res.start: res.end]
            op = self._operators.get(res.entity_type, self._operators.get("DEFAULT"))

            if op.operator_name == "custom":
                gen_fn: Callable[[str], str] = op.params["lambda"]
            else:
                # replace operator
                new_val = op.params.get("new_value", f"<{res.entity_type}>")
                gen_fn = lambda _, v=new_val: v

            fake = self._mapper.get_or_create(res.entity_type, original, gen_fn)
            replacement_map[(res.start, res.end)] = fake
            self._stats.record_detection(
                entity_type=res.entity_type,
                start=res.start,
                end=res.end,
                text=original,
                score=res.score,
                source=source,
            )

        return replacement_map

    # ── Run-level replacement (preserves formatting) ──────────────────────────

    def _apply_replacements_to_runs(
        self,
        para: Paragraph,
        original_full_text: str,
        replacement_map: Dict[Tuple[int, int], str],
    ) -> None:
        """
        Distribute the redacted text back across Word runs, preserving each
        run's individual formatting (bold, italic, font colour, size, etc.).

        Strategy:
          1. Build a character-offset map: which run owns which char offset.
          2. Build the fully redacted text string.
          3. For each run, compute which slice of the redacted text belongs
             to that run (proportional to original char widths), and set only
             that slice. This ensures bold/italic/colour on every run survives.
        """
        if not replacement_map or not para.runs:
            return

        runs = para.runs
        if not runs:
            return

        # Verify runs concatenate to the original text we analysed.
        run_text = "".join(run.text for run in runs)
        if run_text != original_full_text:
            # Structural mismatch (e.g. field code split) — fall back safely.
            para.clear()
            new_run = para.add_run(original_full_text)
            redacted = original_full_text
            for (start, end), fake in sorted(replacement_map.items(), reverse=True):
                redacted = redacted[:start] + fake + redacted[end:]
            new_run.text = redacted
            return

        # Build (run_index, start_offset, end_offset) for each run.
        run_offsets: List[Tuple[int, int, int]] = []
        cursor = 0
        for i, run in enumerate(runs):
            run_offsets.append((i, cursor, cursor + len(run.text)))
            cursor += len(run.text)

        # Apply replacements in reverse order to keep offsets valid.
        # Produce a list of (start, end, fake) tuples sorted reverse.
        replacements = sorted(replacement_map.items(), reverse=True)

        # Build the fully redacted string once.
        redacted = original_full_text
        for (start, end), fake in replacements:
            redacted = redacted[:start] + fake + redacted[end:]

        # Compute the character-length delta introduced by the replacements.
        # Then distribute the redacted text across runs according to original
        # run boundaries, adjusted for cumulative delta.
        #
        # We track a cumulative offset shift as we scan left-to-right.
        # For each run we know its original [run_start, run_end) in the
        # original string. We compute the corresponding slice in the redacted
        # string by accounting for all substitutions that fall before / inside
        # this run.

        # Build a sorted list of substitutions for offset arithmetic.
        subs = sorted([(s, e, fake) for (s, e), fake in replacement_map.items()])

        delta = 0  # cumulative character shift so far
        sub_idx = 0  # pointer into subs

        for run_i, orig_start, orig_end in run_offsets:
            # Advance delta for any subs that are fully BEFORE this run.
            while sub_idx < len(subs) and subs[sub_idx][1] <= orig_start:
                s, e, fake = subs[sub_idx]
                delta += len(fake) - (e - s)
                sub_idx += 1

            redacted_start = orig_start + delta

            # Now figure out where this run ends in the redacted string.
            # We need to account for subs that START inside this run.
            temp_delta = delta
            temp_sub_idx = sub_idx
            inner_orig_cursor = orig_start

            while temp_sub_idx < len(subs) and subs[temp_sub_idx][0] < orig_end:
                s, e, fake = subs[temp_sub_idx]
                if s >= orig_start:  # sub starts inside this run
                    temp_delta += len(fake) - (e - s)
                temp_sub_idx += 1

            redacted_end = orig_end + temp_delta
            runs[run_i].text = redacted[redacted_start:redacted_end]

        logger.debug(
            "Replaced paragraph (%d chars → %d chars, %d runs preserved)",
            len(original_full_text), len(redacted), len(runs),
        )

    # ── Table processing ──────────────────────────────────────────────────────

    def _redact_table(self, table, source_prefix: str) -> None:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_id = cell._tc
                if cell_id in self._seen_table_cells:
                    continue
                self._seen_table_cells.add(cell_id)

                cell_prefix = f"{source_prefix}_row_{row_index}_cell_{cell_index}"
                for para_index, para in enumerate(cell.paragraphs):
                    self._redact_paragraph(
                        para,
                        source=f"{cell_prefix}_para_{para_index}",
                    )
                for nested_index, nested_table in enumerate(cell.tables):
                    self._redact_table(
                        nested_table,
                        source_prefix=f"{cell_prefix}_nested_{nested_index}",
                    )
